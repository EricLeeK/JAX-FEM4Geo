"""
Create Word document from CLAUDE.md
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def add_markdown_to_docx(md_file, output_file):
    """Convert markdown content to Word document"""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    in_code_block = False
    code_lang = None

    while i < len(lines):
        line = lines[i]

        # Handle code blocks
        if line.strip().startswith('```'):
            in_code_block = not in_code_block
            if in_code_block:
                code_lang = line.strip()[3:].strip()
            i += 1
            continue

        if in_code_block:
            # Add code with monospace font
            p = doc.add_paragraph(line.rstrip())
            p.style = 'Normal'
            for run in p.runs:
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0, 0, 128)
            i += 1
            continue

        # Handle headings
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=4)

        # Handle bullet points
        elif line.strip().startswith('- '):
            text = line.strip()[2:]
            # Handle inline code
            text = re.sub(r'`([^`]+)`', r'\1', text)
            p = doc.add_paragraph(text, style='List Bullet')

        # Handle numbered lists
        elif re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            # Handle bold text
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            # Handle inline code
            text = re.sub(r'`([^`]+)`', r'\1', text)
            p = doc.add_paragraph(text, style='List Number')

        # Handle regular paragraphs
        elif line.strip() and not line.startswith('#'):
            text = line.strip()
            # Handle inline code
            parts = re.split(r'(`[^`]+`)', text)
            p = doc.add_paragraph()
            for part in parts:
                if part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(192, 0, 0)
                else:
                    # Handle bold text
                    bold_parts = re.split(r'(\*\*[^*]+\*\*)', part)
                    for bp in bold_parts:
                        if bp.startswith('**') and bp.endswith('**'):
                            run = p.add_run(bp[2:-2])
                            run.bold = True
                        else:
                            p.add_run(bp)

        i += 1

    doc.save(output_file)
    print(f"Word文档已创建: {output_file}")

if __name__ == "__main__":
    add_markdown_to_docx('CLAUDE.md', 'CLAUDE_文档.docx')
